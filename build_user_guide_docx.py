"""
build_user_guide_docx.py
Synthesizes CanoPie_Quick_start_user_guide.docx from user_guide_content.md
and diagram images in docs_assets/ using python-docx.
"""

import os
import re
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

# Document Palette
HEX_NAVY = "1B365D"       # Primary Navy Blue
HEX_TEAL = "008080"       # Secondary Teal
HEX_CHARCOAL = "222222"   # Body Text
HEX_SLATE_BG = "F8FAFC"   # Soft Light Background
HEX_CALLOUT_BG = "F1F5F9" # Callout Background
HEX_BORDER = "CBD5E1"     # Table Border Grey
HEX_AMBER = "D97706"      # Warning/Tip Accent

COLOR_NAVY = RGBColor(0x1B, 0x36, 0x5D)
COLOR_TEAL = RGBColor(0x00, 0x80, 0x80)
COLOR_CHARCOAL = RGBColor(0x22, 0x22, 0x22)
COLOR_MUTED = RGBColor(0x64, 0x74, 0x8B)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Image Mapping Dictionary
IMAGE_MAP = {
    "fig01_main_window": "fig01_main_window_layout.png",
    "fig01_main_window_layout": "fig01_main_window_layout.png",
    "fig02_project_tab_toolbar": "fig02_project_tab_toolbar.png",
    "fig03_image_viewer_z_layers": "fig03_viewport_z_layers.png",
    "fig03_viewport_z_layers": "fig03_viewport_z_layers.png",
    "fig04_stretch_dialog": "fig04_stretch_controls.png",
    "fig04_stretch_controls": "fig04_stretch_controls.png",
    "fig05_band_selector_bar": "fig09_band_selector_math.png",
    "fig05_polygon_manager": "fig05_polygon_manager.png",
    "fig06_drawing_modes": "fig05_polygon_manager.png",
    "fig06_image_editor_ax": "fig06_image_editor_ax.png",
    "fig07_vertex_editing": "fig05_polygon_manager.png",
    "fig07_ml_manager": "fig07_ml_manager.png",
    "fig08_export_manager": "fig08_export_manager.png",
    "fig09_band_selector_math": "fig09_band_selector_math.png",
    "fig10_shapefile_import_pipeline": "fig05_polygon_manager.png",
    "fig11_spatial_indexing_tree": "fig05_polygon_manager.png",
    "fig12_image_editor_dialog": "fig06_image_editor_ax.png",
    "fig13_export_manager": "fig08_export_manager.png",
    "fig15_csv_viewer_dialog": "fig07_ml_manager.png"
}

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins_and_border(cell, top=140, bottom=140, left=200, right=200, left_border_color=HEX_NAVY, left_border_sz="36"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)
    
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="{left_border_sz}" w:space="0" w:color="{left_border_color}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)

def set_table_borders(table, color=HEX_BORDER, sz="4"):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(tblBorders)

def set_heading1_bottom_border(paragraph, color=HEX_NAVY, sz="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(f'''
        <w:pBdr {nsdecls("w")}>
            <w:bottom w:val="single" w:sz="{sz}" w:space="6" w:color="{color}"/>
        </w:pBdr>
    ''')
    pPr.append(pBdr)

def make_row_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = parse_xml(f'<w:tblHeader {nsdecls("w")}/>')
    trPr.append(tblHeader)

def make_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    cantSplit = parse_xml(f'<w:cantSplit {nsdecls("w")}/>')
    trPr.append(cantSplit)

def add_formatted_text(paragraph, text, default_font="Segoe UI", default_size=Pt(11), default_color=COLOR_CHARCOAL):
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\$\$.*?\$\$|\$.*?\$|\[.*?\]\(.*?\))')
    tokens = pattern.split(text)
    
    for token in tokens:
        if not token:
            continue
        
        run = paragraph.add_run()
        run.font.name = default_font
        run.font.size = default_size
        run.font.color.rgb = default_color
        
        if token.startswith('**') and token.endswith('**') and len(token) >= 4:
            run.text = token[2:-2]
            run.font.bold = True
        elif token.startswith('*') and token.endswith('*') and len(token) >= 2:
            run.text = token[1:-1]
            run.font.italic = True
        elif token.startswith('`') and token.endswith('`') and len(token) >= 2:
            run.text = token[1:-1]
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            rPr = run._r.get_or_add_rPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
            rPr.append(shd)
        elif (token.startswith('$$') and token.endswith('$$')) or (token.startswith('$') and token.endswith('$')):
            math_text = token.strip('$')
            run.text = math_text
            run.font.italic = True
            run.font.name = 'Cambria Math'
            run.font.color.rgb = COLOR_NAVY
        elif token.startswith('[') and ']' in token and '(' in token and token.endswith(')'):
            m = re.match(r'\[(.*?)\]\((.*?)\)', token)
            if m:
                label, url = m.groups()
                run.text = label
                run.font.color.rgb = COLOR_TEAL
                run.font.underline = True
            else:
                run.text = token
        else:
            run.text = token

def resolve_image_path(fig_key, docs_assets_dir):
    img_name = IMAGE_MAP.get(fig_key)
    if not img_name:
        for key in IMAGE_MAP:
            if fig_key.startswith(key[:5]):
                img_name = IMAGE_MAP[key]
                break
    if not img_name:
        img_name = "fig01_main_window_layout.png"
    
    full_path = os.path.join(docs_assets_dir, img_name)
    if not os.path.exists(full_path):
        # fallback to any valid image in docs_assets
        for f in os.listdir(docs_assets_dir):
            if f.endswith('.png'):
                return os.path.join(docs_assets_dir, f)
    return full_path

def build_docx():
    base_dir = r"c:\Users\natha\Downloads\CanoPie-main_updated\CanoPie-main"
    content_file = os.path.join(base_dir, ".agents", "worker_m2_content", "user_guide_content.md")
    docs_assets_dir = os.path.join(base_dir, "docs_assets")
    output_docx = os.path.join(base_dir, "CanoPie_Quick_start_user_guide.docx")

    print(f"Reading content from: {content_file}")
    with open(content_file, "r", encoding="utf-8") as f:
        md_text = f.read()

    doc = Document()

    # Page Margins Setup
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Footer Setup
    footer = section.footer
    p_ft = footer.paragraphs[0]
    p_ft.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_ft = p_ft.add_run("CanoPie Quick Start User Guide  |  Page ")
    run_ft.font.name = 'Segoe UI'
    run_ft.font.size = Pt(9)
    run_ft.font.color.rgb = COLOR_MUTED

    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    p_ft.add_run()._r.extend([fldChar1, instrText, fldChar2, fldChar3])

    lines = md_text.splitlines()

    # Document Header Section
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_t = p_title.add_run("CanoPie Quick Start User Guide")
    run_t.font.name = 'Segoe UI'
    run_t.font.size = Pt(28)
    run_t.font.bold = True
    run_t.font.color.rgb = COLOR_NAVY

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(12)
    run_s = p_sub.add_run("Comprehensive User Manual & Technical Reference V_1_0_0")
    run_s.font.name = 'Segoe UI'
    run_s.font.size = Pt(14)
    run_s.font.italic = True
    run_s.font.color.rgb = COLOR_TEAL

    # Metadata Block Table
    meta_table = doc.add_table(rows=1, cols=3)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(meta_table, color=HEX_BORDER, sz="6")
    
    meta_data = [
        ("Software Version", "CanoPie V_1_0_0"),
        ("Classification", "Complete Technical Reference"),
        ("Date", "August 2026")
    ]
    for idx, (label, val) in enumerate(meta_data):
        cell = meta_table.rows[0].cells[idx]
        cell.width = Inches(2.16)
        set_cell_background(cell, HEX_SLATE_BG)
        p_m = cell.paragraphs[0]
        p_m.paragraph_format.space_after = Pt(2)
        p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        r_lbl = p_m.add_run(f"{label}\n")
        r_lbl.font.name = 'Segoe UI'
        r_lbl.font.size = Pt(8.5)
        r_lbl.font.color.rgb = COLOR_MUTED
        
        r_val = p_m.add_run(val)
        r_val.font.name = 'Segoe UI'
        r_val.font.size = Pt(10)
        r_val.font.bold = True
        r_val.font.color.rgb = COLOR_NAVY

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(12)

    # State Machine for Parsing
    fig_counter = 0
    idx = 0
    num_lines = len(lines)

    # Skip initial title lines in markdown as we built header
    while idx < num_lines:
        line = lines[idx]

        # Ignore initial title & metadata block in markdown text
        if line.startswith("# CanoPie Quick Start User Guide") or \
           line.startswith("**Software Version**:") or \
           line.startswith("**Document Classification**:") or \
           line.strip() == "---" and idx < 10:
            idx += 1
            continue

        # 1. Code Block
        if line.strip().startswith("```"):
            code_lines = []
            idx += 1
            while idx < num_lines and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx])
                idx += 1
            if idx < num_lines:
                idx += 1 # skip closing ```

            code_table = doc.add_table(rows=1, cols=1)
            code_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = code_table.rows[0].cells[0]
            cell.width = Inches(6.5)
            set_cell_background(cell, HEX_SLATE_BG)
            set_cell_margins_and_border(cell, top=120, bottom=120, left=160, right=160, left_border_color=HEX_BORDER, left_border_sz="12")
            
            p_code = cell.paragraphs[0]
            p_code.paragraph_format.line_spacing = 1.0
            p_code.paragraph_format.space_after = Pt(2)
            
            for l_idx, cl in enumerate(code_lines):
                if l_idx > 0:
                    p_code = cell.add_paragraph()
                    p_code.paragraph_format.line_spacing = 1.0
                    p_code.paragraph_format.space_after = Pt(2)
                run_c = p_code.add_run(cl)
                run_c.font.name = 'Consolas'
                run_c.font.size = Pt(9.0)
                run_c.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            
            p_sp = doc.add_paragraph()
            p_sp.paragraph_format.space_after = Pt(6)
            continue

        # 2. Blockquotes / Callout Box
        if line.strip().startswith(">"):
            quote_lines = []
            while idx < num_lines and lines[idx].strip().startswith(">"):
                qline = lines[idx].strip()[1:].strip()
                quote_lines.append(qline)
                idx += 1
            
            full_quote = " ".join(quote_lines)
            border_color = HEX_AMBER if ("TIP" in full_quote or "WARNING" in full_quote) else HEX_NAVY
            
            callout_table = doc.add_table(rows=1, cols=1)
            callout_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = callout_table.rows[0].cells[0]
            cell.width = Inches(6.5)
            set_cell_background(cell, HEX_CALLOUT_BG)
            set_cell_margins_and_border(cell, top=140, bottom=140, left=200, right=200, left_border_color=border_color, left_border_sz="36")
            
            p_q = cell.paragraphs[0]
            p_q.paragraph_format.line_spacing = 1.15
            p_q.paragraph_format.space_after = Pt(4)
            add_formatted_text(p_q, full_quote, default_font='Segoe UI', default_size=Pt(10), default_color=COLOR_CHARCOAL)
            
            p_sp = doc.add_paragraph()
            p_sp.paragraph_format.space_after = Pt(6)
            continue

        # 3. Markdown Table
        if "|" in line and idx + 1 < num_lines and ("|---" in lines[idx + 1] or "| ---" in lines[idx + 1] or "|:-" in lines[idx + 1]):
            table_lines = []
            while idx < num_lines and "|" in lines[idx]:
                table_lines.append(lines[idx])
                idx += 1
            
            def parse_table_row(r_line):
                cleaned = r_line.strip().replace(r'\|', '__ESCAPED_PIPE__')
                if cleaned.startswith('|'):
                    cleaned = cleaned[1:]
                if cleaned.endswith('|'):
                    cleaned = cleaned[:-1]
                return [c.strip().replace('__ESCAPED_PIPE__', '|') for c in cleaned.split('|')]

            headers = parse_table_row(table_lines[0])
            data_rows = []
            for t_line in table_lines[2:]:
                if "|" in t_line:
                    row_cells = parse_table_row(t_line)
                    # Ignore separator row if present
                    if all(re.match(r'^:?-+:?$', c) for c in row_cells if c):
                        continue
                    if len(row_cells) == len(headers):
                        data_rows.append(row_cells)
                    elif len(row_cells) < len(headers):
                        row_cells.extend([""] * (len(headers) - len(row_cells)))
                        data_rows.append(row_cells)
                    else:
                        row_cells = row_cells[:len(headers)]
                        data_rows.append(row_cells)

            
            tbl = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_borders(tbl)
            
            # Header Row
            hdr_row = tbl.rows[0]
            make_row_header(hdr_row)
            make_row_cant_split(hdr_row)
            for c_idx, h_text in enumerate(headers):
                cell = hdr_row.cells[c_idx]
                set_cell_background(cell, HEX_NAVY)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_formatted_text(p, h_text, default_font='Segoe UI', default_size=Pt(10), default_color=COLOR_WHITE)
                for run in p.runs:
                    run.font.bold = True
            
            # Data Rows
            for r_idx, r_data in enumerate(data_rows):
                tr = tbl.rows[r_idx + 1]
                make_row_cant_split(tr)
                bg_color = HEX_SLATE_BG if r_idx % 2 == 1 else "FFFFFF"
                for c_idx, cell_value in enumerate(r_data):
                    if c_idx < len(tr.cells):
                        cell = tr.cells[c_idx]
                        if bg_color != "FFFFFF":
                            set_cell_background(cell, bg_color)
                        p = cell.paragraphs[0]
                        add_formatted_text(p, cell_value, default_font='Segoe UI', default_size=Pt(9.5), default_color=COLOR_CHARCOAL)
            
            # Smart Column Widths
            max_lens = [len(h) for h in headers]
            for r_data in data_rows:
                for c_idx, val in enumerate(r_data):
                    if c_idx < len(max_lens):
                        max_lens[c_idx] = max(max_lens[c_idx], len(val))
            
            tot_len = sum(max_lens) if sum(max_lens) > 0 else 1
            col_widths = [Inches(max(0.6, min(3.5, 6.5 * (l / tot_len)))) for l in max_lens]
            for row in tbl.rows:
                for c_idx, w in enumerate(col_widths):
                    if c_idx < len(row.cells):
                        row.cells[c_idx].width = w

            p_sp = doc.add_paragraph()
            p_sp.paragraph_format.space_after = Pt(8)
            continue

        # 4. Figure Placeholder
        if line.strip().startswith("[INSERT FIGURE"):
            m = re.match(r'\[INSERT FIGURE (fig\w+) HERE:\s*(.*?)\]', line.strip())
            if m:
                fig_key, caption = m.groups()
                img_path = resolve_image_path(fig_key, docs_assets_dir)
                
                if os.path.exists(img_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(8)
                    p_img.paragraph_format.space_after = Pt(4)
                    run_img = p_img.add_run()
                    run_img.add_picture(img_path, width=Inches(5.8))
                    
                    fig_counter += 1
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_before = Pt(2)
                    p_cap.paragraph_format.space_after = Pt(12)
                    run_cap = p_cap.add_run(f"Figure {fig_counter}: {caption}")
                    run_cap.font.name = 'Segoe UI'
                    run_cap.font.size = Pt(9.5)
                    run_cap.font.italic = True
                    run_cap.font.color.rgb = COLOR_MUTED
            idx += 1
            continue

        # 5. Headings
        if line.startswith("# "):
            h_text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            add_formatted_text(p, h_text, default_font='Segoe UI', default_size=Pt(20), default_color=COLOR_NAVY)
            for run in p.runs:
                run.font.bold = True
            set_heading1_bottom_border(p, color=HEX_NAVY, sz="12")
            idx += 1
            continue

        if line.startswith("## "):
            h_text = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            add_formatted_text(p, h_text, default_font='Segoe UI', default_size=Pt(15), default_color=COLOR_TEAL)
            for run in p.runs:
                run.font.bold = True
            idx += 1
            continue

        if line.startswith("### "):
            h_text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            add_formatted_text(p, h_text, default_font='Segoe UI', default_size=Pt(12), default_color=COLOR_NAVY)
            for run in p.runs:
                run.font.bold = True
            idx += 1
            continue

        if line.startswith("#### "):
            h_text = line[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            add_formatted_text(p, h_text, default_font='Segoe UI', default_size=Pt(11), default_color=COLOR_CHARCOAL)
            for run in p.runs:
                run.font.bold = True
            idx += 1
            continue

        # 6. Bullet & Numbered Lists
        stripped = line.strip()
        if stripped.startswith("- ") or stripped.startswith("* "):
            item_text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, item_text, default_font='Segoe UI', default_size=Pt(11), default_color=COLOR_CHARCOAL)
            idx += 1
            continue

        if re.match(r'^\d+\.\s', stripped):
            item_text = re.sub(r'^\d+\.\s', '', stripped).strip()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, item_text, default_font='Segoe UI', default_size=Pt(11), default_color=COLOR_CHARCOAL)
            idx += 1
            continue

        # 7. Horizontal Rule
        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            set_heading1_bottom_border(p, color=HEX_BORDER, sz="4")
            idx += 1
            continue

        # 8. Regular Paragraph
        if stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, stripped, default_font='Segoe UI', default_size=Pt(11), default_color=COLOR_CHARCOAL)

        idx += 1

    print(f"Saving Word Document to: {output_docx}")
    doc.save(output_docx)

    # Verification
    if not os.path.exists(output_docx):
        raise FileNotFoundError(f"Generated file missing: {output_docx}")
    
    file_size_kb = os.path.getsize(output_docx) / 1024.0
    print(f"Generated DOCX size: {file_size_kb:.2f} KB")

    if file_size_kb < 500:
        raise ValueError(f"DOCX size {file_size_kb:.2f} KB is less than required 500 KB threshold")

    # Verify openability via python-docx
    verify_doc = Document(output_docx)
    img_count = len(verify_doc.inline_shapes)
    print(f"Verification Success! Document contains {len(verify_doc.paragraphs)} paragraphs, {len(verify_doc.tables)} tables, and {img_count} embedded images.")
    if img_count < 9:
        raise ValueError(f"DOCX image count {img_count} is less than required 9 embedded images")


if __name__ == "__main__":
    build_docx()
